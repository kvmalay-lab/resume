import os
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

from fpdf import FPDF

def extract_text_from_pdf(file_path: str) -> str:
    if not PdfReader:
        return "Error: pypdf is not installed."
    try:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n".join(text)
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(file_path: str) -> str:
    if not docx:
        return "Error: python-docx is not installed."
    try:
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def extract_text(file_path: str) -> str:
    """Helper to route file parsing based on extension."""
    _, ext = os.path.splitext(file_path.lower())
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            return f"Error reading text file: {str(e)}"
    else:
        return "Error: Unsupported file format. Please upload a PDF, DOCX, or TXT file."


def export_to_pdf(text: str, output_path: str):
    """Exports text to a PDF file."""
    try:
        pdf = FPDF()
        pdf.add_page()
        # Fallback to standard core font, encoding might strip advanced unicode emojis
        pdf.set_font("Arial", size=11)
        encoded_text = text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 5, encoded_text)
        pdf.output(output_path)
        return output_path
    except Exception as e:
        print(f"Error exporting to PDF: {str(e)}")
        return None

def export_to_docx(text: str, output_path: str):
    """Exports text to a DOCX file."""
    if not docx:
        print("python-docx not installed.")
        return None
    try:
        doc = docx.Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Error exporting to DOCX: {str(e)}")
        return None
